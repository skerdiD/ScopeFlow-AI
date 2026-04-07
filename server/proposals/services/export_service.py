from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime, timezone
from html import escape
from io import BytesIO

from docx import Document
from docx.shared import Pt
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer


@dataclass
class MilestoneItem:
    title: str
    description: str


@dataclass
class ExportSections:
    summary: list[str]
    scope_of_work: list[str]
    deliverables: list[str]
    milestones: list[MilestoneItem]
    risks: list[str]


def _safe_filename_part(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip().lower()).strip("-")
    return cleaned or fallback


def build_export_filename(project_name: str, source_label: str, file_ext: str) -> str:
    name_part = _safe_filename_part(project_name, "proposal")
    source_part = _safe_filename_part(source_label, "current")
    date_part = datetime.now(timezone.utc).strftime("%Y%m%d")
    return f"{name_part}-{source_part}-{date_part}.{file_ext}"


def _to_paragraphs(value: str) -> list[str]:
    if not value.strip():
        return ["Not provided."]

    paragraphs: list[str] = []
    current_lines: list[str] = []
    for raw_line in value.splitlines():
        line = raw_line.strip()
        if not line:
            if current_lines:
                paragraphs.append(" ".join(current_lines))
                current_lines = []
            continue
        current_lines.append(line)

    if current_lines:
        paragraphs.append(" ".join(current_lines))

    return paragraphs or ["Not provided."]


def _to_list_items(value: str) -> list[str]:
    if not value.strip():
        return ["Not provided."]

    items: list[str] = []
    for raw_line in value.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^[-*•]\s*", "", line)
        line = re.sub(r"^\d+[\.\)]\s*", "", line)
        cleaned = line.strip()
        if cleaned:
            items.append(cleaned)

    return items or ["Not provided."]


def _to_milestones(value: str) -> list[MilestoneItem]:
    if not value.strip():
        return [MilestoneItem(title="Milestone", description="Not provided.")]

    milestones: list[MilestoneItem] = []
    for raw_line in value.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^[-*•]\s*", "", line)
        parts = line.split(":", 1)
        if len(parts) == 2:
            title = parts[0].strip() or "Milestone"
            description = parts[1].strip() or "Not provided."
        else:
            title = "Milestone"
            description = line.strip() or "Not provided."
        milestones.append(MilestoneItem(title=title, description=description))

    return milestones or [MilestoneItem(title="Milestone", description="Not provided.")]


def build_export_sections(
    *,
    summary: str,
    scope: str,
    deliverables: str,
    milestones: str,
    risks: str,
) -> ExportSections:
    return ExportSections(
        summary=_to_paragraphs(summary),
        scope_of_work=_to_list_items(scope),
        deliverables=_to_list_items(deliverables),
        milestones=_to_milestones(milestones),
        risks=_to_list_items(risks),
    )


def _append_section_header(story: list, heading_style: ParagraphStyle, title: str):
    story.append(Spacer(1, 14))
    story.append(Paragraph(escape(title), heading_style))
    story.append(Spacer(1, 6))


def _append_paragraphs(story: list, body_style: ParagraphStyle, paragraphs: list[str]):
    for paragraph in paragraphs:
        story.append(Paragraph(escape(paragraph), body_style))
        story.append(Spacer(1, 4))


def _append_bullets(story: list, bullet_style: ParagraphStyle, items: list[str]):
    list_items = [ListItem(Paragraph(escape(item), bullet_style)) for item in items]
    story.append(ListFlowable(list_items, bulletType="bullet", leftIndent=12))
    story.append(Spacer(1, 4))


def generate_pdf_export(
    *,
    project_name: str,
    client_name: str,
    project_type: str,
    budget: str,
    timeline: str,
    source_label: str,
    sections: ExportSections,
) -> bytes:
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        leftMargin=50,
        rightMargin=50,
        topMargin=50,
        bottomMargin=50,
        title=project_name or "Proposal",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ScopeFlowTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=HexColor("#0f172a"),
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "ScopeFlowHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=HexColor("#111827"),
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ScopeFlowBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=HexColor("#1f2937"),
    )
    bullet_style = ParagraphStyle(
        "ScopeFlowBullet",
        parent=body_style,
        leftIndent=8,
    )
    metadata_style = ParagraphStyle(
        "ScopeFlowMetadata",
        parent=body_style,
        textColor=HexColor("#374151"),
    )

    story: list = []
    story.append(Paragraph(escape(project_name or "Proposal"), title_style))
    story.append(Paragraph(f"<b>Client:</b> {escape(client_name or '-')}", metadata_style))
    story.append(Paragraph(f"<b>Project Type:</b> {escape(project_type or '-')}", metadata_style))
    story.append(Paragraph(f"<b>Budget:</b> {escape(budget or '-')}", metadata_style))
    story.append(Paragraph(f"<b>Timeline:</b> {escape(timeline or '-')}", metadata_style))
    story.append(Paragraph(f"<b>Export Source:</b> {escape(source_label)}", metadata_style))
    story.append(Spacer(1, 10))

    _append_section_header(story, heading_style, "Summary")
    _append_paragraphs(story, body_style, sections.summary)

    _append_section_header(story, heading_style, "Scope of Work")
    _append_bullets(story, bullet_style, sections.scope_of_work)

    _append_section_header(story, heading_style, "Deliverables")
    _append_bullets(story, bullet_style, sections.deliverables)

    _append_section_header(story, heading_style, "Milestones")
    _append_bullets(
        story,
        bullet_style,
        [f"{milestone.title}: {milestone.description}" for milestone in sections.milestones],
    )

    _append_section_header(story, heading_style, "Risks")
    _append_bullets(story, bullet_style, sections.risks)

    story.append(Spacer(1, 14))
    story.append(
        Paragraph(
            escape(f"Generated by ScopeFlow AI on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"),
            ParagraphStyle(
                "ScopeFlowFooter",
                parent=body_style,
                fontSize=9,
                textColor=HexColor("#6b7280"),
            ),
        )
    )

    document.build(story)
    return buffer.getvalue()


def generate_docx_export(
    *,
    project_name: str,
    client_name: str,
    project_type: str,
    budget: str,
    timeline: str,
    source_label: str,
    sections: ExportSections,
) -> bytes:
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    document.add_heading(project_name or "Proposal", level=0)

    metadata_lines = [
        ("Client", client_name or "-"),
        ("Project Type", project_type or "-"),
        ("Budget", budget or "-"),
        ("Timeline", timeline or "-"),
        ("Export Source", source_label),
    ]
    for label, value in metadata_lines:
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    document.add_paragraph("")

    document.add_heading("Summary", level=1)
    for paragraph in sections.summary:
        document.add_paragraph(paragraph)

    document.add_heading("Scope of Work", level=1)
    for item in sections.scope_of_work:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Deliverables", level=1)
    for item in sections.deliverables:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Milestones", level=1)
    for item in sections.milestones:
        milestone_paragraph = document.add_paragraph(style="List Bullet")
        milestone_paragraph.add_run(f"{item.title}: ").bold = True
        milestone_paragraph.add_run(item.description)

    document.add_heading("Risks", level=1)
    for item in sections.risks:
        document.add_paragraph(item, style="List Bullet")

    document.add_paragraph("")
    document.add_paragraph(f"Generated by ScopeFlow AI on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    output = BytesIO()
    document.save(output)
    return output.getvalue()
